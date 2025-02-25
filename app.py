import streamlit as st
import pandas as pd
from PIL import Image
from docx import Document
import PyPDF2
import io
import plotly.express as px
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np

# Set page config
st.set_page_config(
    page_title="Interactive Document & Image Processor",
    page_icon="📄",
    layout="wide"
)

# Enhanced custom CSS with more animations and colors
st.markdown("""
    <style>
    /* Main container styling */
    .main {
        padding: 2rem;
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
    }

    /* Title animations and styling */
    .stTitle {
        background: linear-gradient(120deg, #155799, #159957);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 3.5rem !important;
        font-weight: 700 !important;
        animation: titleAnimation 2s ease-in-out;
    }

    /* Animations */
    @keyframes titleAnimation {
        0% { transform: translateY(-20px); opacity: 0; }
        100% { transform: translateY(0); opacity: 1; }
    }

    @keyframes floatingAnimation {
        0% { transform: translateY(0px); }
        50% { transform: translateY(-10px); }
        100% { transform: translateY(0px); }
    }

    @keyframes glowingButton {
        0% { box-shadow: 0 0 5px #4CAF50; }
        50% { box-shadow: 0 0 20px #4CAF50; }
        100% { box-shadow: 0 0 5px #4CAF50; }
    }

    /* Button styling */
    .stButton>button {
        background: linear-gradient(45deg, #4CAF50, #45a049);
        border: none;
        color: white;
        padding: 0.5rem 1rem;
        border-radius: 10px;
        font-weight: 600;
        transition: all 0.3s ease;
        animation: glowingButton 2s infinite;
    }

    .stButton>button:hover {
        transform: scale(1.05);
        box-shadow: 0 10px 20px rgba(0,0,0,0.2);
    }

    /* Metric containers */
    .css-1r6slb0 {
        background: rgba(255, 255, 255, 0.1);
        border-radius: 15px;
        padding: 1rem;
        backdrop-filter: blur(10px);
        border: 1px solid rgba(255, 255, 255, 0.2);
        transition: transform 0.3s ease;
    }

    .css-1r6slb0:hover {
        transform: translateY(-5px);
    }

    /* Tab styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 2rem;
        background-color: transparent;
    }

    .stTabs [data-baseweb="tab"] {
        height: 50px;
        background-color: rgba(255, 255, 255, 0.1);
        border-radius: 10px;
        color: #31333F;
        font-weight: 600;
        transition: all 0.3s ease;
    }

    .stTabs [data-baseweb="tab"]:hover {
        background-color: rgba(255, 255, 255, 0.2);
        color: #4CAF50;
    }

    /* Sidebar styling */
    .css-1d391kg {
        background: linear-gradient(180deg, #2c3e50 0%, #3498db 100%);
    }

    /* File uploader styling */
    .uploadedFile {
        background: rgba(255, 255, 255, 0.1);
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
        border: 2px dashed #4CAF50;
    }

    /* Progress bar animation */
    @keyframes progressAnimation {
        0% { width: 0%; }
        100% { width: 100%; }
    }

    /* Slider styling */
    .stSlider {
        padding: 1rem 0;
    }

    .stSlider > div > div {
        background-color: #4CAF50;
    }

    /* Success message styling */
    .stSuccess {
        animation: fadeIn 0.5s ease-in;
        background: linear-gradient(45deg, #4CAF50, #45a049);
        border-radius: 10px;
        border: none;
    }

    /* Error message styling */
    .stError {
        animation: fadeIn 0.5s ease-in;
        background: linear-gradient(45deg, #ff4b4b, #ff0000);
        border-radius: 10px;
        border: none;
    }

    /* Footer styling */
    .footer {
        text-align: center;
        padding: 2rem;
        background: linear-gradient(45deg, #2c3e50, #3498db);
        color: white;
        border-radius: 10px;
        margin-top: 2rem;
    }

    .footer a {
        color: #4CAF50;
        text-decoration: none;
        transition: color 0.3s ease;
    }

    .footer a:hover {
        color: #45a049;
    }
    </style>
""", unsafe_allow_html=True)

# Sidebar configuration
st.sidebar.title("⚙️ Settings")
theme = st.sidebar.selectbox("Choose Theme", ["Light", "Dark"])
if theme == "Dark":
    plt.style.use("dark_background")

# Title and description with animation
st.title("📄 Interactive Document & Image Processor")
st.markdown("### Transform your documents and images with powerful tools! 🚀")

# Create tabs for different file types
tab1, tab2, tab3 = st.tabs(["📝 DOCX", "📋 PDF", "🖼️ Images"])

with tab1:
    st.header("DOCX Processing")
    docx_file = st.file_uploader("Upload a DOCX file", type=['docx'])
    
    if docx_file:
        doc = Document(docx_file)
        text_content = []
        
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        
        # Interactive Document Statistics with Animation
        st.subheader("📊 Document Statistics")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Paragraphs", len(doc.paragraphs), delta="📝")
        with col2:
            words = sum(len(para.split()) for para in text_content)
            st.metric("Words", words, delta="📚")
        with col3:
            characters = sum(len(para) for para in text_content)
            st.metric("Characters", characters, delta="📖")
        with col4:
            avg_words_per_para = round(words / len(doc.paragraphs), 2)
            st.metric("Avg Words/Paragraph", avg_words_per_para, delta="📊")
        
        # Text Analysis Features
        st.subheader("🔍 Text Analysis")
        search_term = st.text_input("Search for word or phrase:")
        if search_term:
            occurrences = sum(text.lower().count(search_term.lower()) for text in text_content)
            st.success(f"Found {occurrences} occurrences of '{search_term}'")
        
        # Word Cloud Option
        if st.checkbox("Show Word Frequency Analysis"):
            words_list = " ".join(text_content).split()
            word_freq = pd.Series(words_list).value_counts().head(10)
            fig = px.bar(x=word_freq.index, y=word_freq.values,
                        title="Top 10 Most Common Words",
                        labels={'x': 'Words', 'y': 'Frequency'})
            st.plotly_chart(fig)
        
        # Display content with formatting options
        st.subheader("📄 Document Content")
        display_mode = st.radio("Display Mode:", ["Full Text", "Paragraph by Paragraph"])
        
        if display_mode == "Full Text":
            st.text_area("Content", "\n".join(text_content), height=300)
        else:
            para_num = st.slider("Select Paragraph", 1, len(doc.paragraphs))
            st.text_area(f"Paragraph {para_num}", doc.paragraphs[para_num-1].text)

with tab2:
    st.header("PDF Processing")
    pdf_file = st.file_uploader("Upload a PDF file", type=['pdf'])
    
    if pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_content = []
        
        for page in pdf_reader.pages:
            text_content.append(page.extract_text())
        
        # Enhanced PDF Statistics
        st.subheader("📊 PDF Statistics")
        col1, col2, col3 = st.columns(3)
        
        total_text = "".join(text_content)
        with col1:
            st.metric("Pages", len(pdf_reader.pages), delta="📄")
        with col2:
            word_count = len(total_text.split())
            st.metric("Words", word_count, delta="📚")
        with col3:
            avg_words_per_page = round(word_count / len(pdf_reader.pages), 2)
            st.metric("Avg Words/Page", avg_words_per_page, delta="📊")
        
        # PDF Analysis Tools
        st.subheader("🔍 PDF Analysis")
        search_pdf = st.text_input("Search PDF content:")
        if search_pdf:
            page_matches = []
            for i, page in enumerate(text_content):
                if search_pdf.lower() in page.lower():
                    page_matches.append(i + 1)
            if page_matches:
                st.success(f"Found '{search_pdf}' on pages: {', '.join(map(str, page_matches))}")
            else:
                st.warning(f"No matches found for '{search_pdf}'")
        
        # Interactive Page Viewer
        st.subheader("📄 Page Viewer")
        col1, col2 = st.columns([1, 3])
        with col1:
            page_num = st.number_input("Go to page", min_value=1, 
                                     max_value=len(pdf_reader.pages), value=1)
        with col2:
            st.slider("Navigate Pages", 1, len(pdf_reader.pages), 
                     value=page_num, key="page_slider")
        
        st.text_area("Page Content", text_content[page_num - 1], height=300)

with tab3:
    st.header("Image Processing")
    image_file = st.file_uploader("Upload an image", type=['png', 'jpg', 'jpeg'])
    
    if image_file:
        image = Image.open(image_file)
        
        # Enhanced Image Information
        st.subheader("📊 Image Information")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Width", f"{image.size[0]}px", delta="📏")
        with col2:
            st.metric("Height", f"{image.size[1]}px", delta="📏")
        with col3:
            st.metric("Format", image.format, delta="🖼️")
        with col4:
            aspect_ratio = round(image.size[0] / image.size[1], 2)
            st.metric("Aspect Ratio", aspect_ratio, delta="📐")
        
        # Advanced Image Processing Options
        st.subheader("🎨 Image Processing")
        
        # Image display with zoom
        zoom_level = st.slider("Zoom Level", 0.5, 2.0, 1.0, 0.1)
        new_size = (int(image.size[0] * zoom_level), int(image.size[1] * zoom_level))
        resized_image = image.resize(new_size)
        st.image(resized_image, use_column_width=True)
        
        # Processing options
        col1, col2 = st.columns(2)
        with col1:
            processing_option = st.selectbox(
                "Select processing option",
                ["Original", "Grayscale", "Color Histogram", "Edge Detection", "Blur"]
            )
        
        with col2:
            if processing_option in ["Blur", "Edge Detection"]:
                intensity = st.slider("Effect Intensity", 0, 10, 5)
        
        # Apply selected effect
        if processing_option == "Grayscale":
            processed_image = image.convert('L')
            st.image(processed_image, caption="Grayscale version", use_column_width=True)
        
        elif processing_option == "Color Histogram":
            fig = plt.figure(figsize=(10, 4))
            if image.mode == 'RGB':
                colors = ('red', 'green', 'blue')
                for i, color in enumerate(colors):
                    hist = image.histogram()[i*256:(i+1)*256]
                    plt.plot(range(256), hist, color=color, alpha=0.7, label=color.capitalize())
                plt.title("RGB Color Distribution")
                plt.xlabel("Color Intensity")
                plt.ylabel("Pixel Count")
                plt.legend()
                st.pyplot(fig)
        
        elif processing_option == "Edge Detection":
            from PIL import ImageFilter
            processed_image = image.convert('L').filter(ImageFilter.FIND_EDGES)
            st.image(processed_image, caption="Edge Detection", use_column_width=True)
        
        elif processing_option == "Blur":
            processed_image = image.filter(ImageFilter.GaussianBlur(intensity))
            st.image(processed_image, caption="Blurred version", use_column_width=True)

        # Image Download Option
        if processing_option != "Original" and processing_option != "Color Histogram":
            buf = io.BytesIO()
            if 'processed_image' in locals():
                processed_image.save(buf, format=image.format)
                st.download_button(
                    label="Download Processed Image",
                    data=buf.getvalue(),
                    file_name=f"processed_{image_file.name}",
                    mime=f"image/{image.format.lower()}"
                )

# Update the footer with more styling
st.markdown("---")
st.markdown("""
    <div class='footer'>
        <h3 style='animation: floatingAnimation 3s infinite ease-in-out'>Rate Your Experience</h3>
        <div style='margin: 2rem 0;'>
""", unsafe_allow_html=True)

feedback = st.slider("How would you rate this app? 🌟", 1, 5, 3)
if feedback >= 4:
    st.success("Thanks for the positive feedback! 🎉 We're glad you enjoyed using our app!")
elif feedback <= 2:
    st.error("We appreciate your feedback! 🔨 We'll work hard to improve your experience!")

st.markdown("""
        </div>
        <div style='margin-top: 2rem;'>
            <p style='font-size: 1.2rem;'>Created with ❤️ using Streamlit</p>
            <div style='display: flex; justify-content: center; gap: 2rem; margin-top: 1rem;'>
                <a href='#' target='_blank'>📚 Documentation</a>
                <a href='#' target='_blank'>🐛 Report Issues</a>
                <a href='#' target='_blank'>💡 Suggestions</a>
            </div>
        </div>
    </div>
""", unsafe_allow_html=True)
